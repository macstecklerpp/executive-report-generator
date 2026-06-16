"""
================================================================================
  PromptPath Executive Summary Performance Report — Version 1
================================================================================
  REQUIREMENTS:  pip install python-docx openpyxl

  HOW TO RUN (CLI)
  1. Fill in the CONFIG block below.
  2. python3 promptpath_exec_report_v1.py
  3. Collect your .docx from OUTPUT_PATH and the companion *_audit.xlsx (number reconciliation).

  PROGRAMMATIC USE
  from promptpath_exec_report_v1 import ReportConfig, generate_report
  generate_report(ReportConfig(...))

  FILES NEEDED
  - Inbound leaderboard CSV    (PromptPath > Leaderboard > Inbound > Export)
  - Outbound leaderboard CSV   (PromptPath > Leaderboard > Outbound > Export)
  - Department calls file      (tab-separated: organization_name, dealer_name,
                                category, calls)
  - PromptPath logo PNG

  NOTES
  - IB_ONLY=True suppresses all outbound metrics and the Sales Dials column.
  - STORE_FILTER limits to stores whose name contains any of the given substrings (or None).
  - Dealer names in the dept file must match the Dealerships column in the CSV.
  - Inbound Opportunities and IB appointment set rate use the unique-customer column when present
    (see IB_UNIQUE_OPP_COLUMNS); otherwise Connected. Outbound Opportunities use OB_UNIQUE_OPP_COLUMNS similarly.
    Outbound connect-count (connect rate numerator, OB appt set rate denominator) uses the first matching column in
    OB_CONNECTED_COLUMN_CANDIDATES (preferring "Connected" when both Connected and "Unique Connected" exist).
  - Optional inbound CSV column Soft Appt for hard-percent denominator; otherwise Soft = Total Appts minus Hard Appt.
================================================================================
"""

from __future__ import annotations

import csv
import re
from dataclasses import dataclass, field, replace
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Union

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════════════════════════
#  CONFIG (CLI defaults)
# ══════════════════════════════════════════════════════════════════════════════
GROUP_NAME = "Dealer Group Name"
PERIOD_START = "2026-05-01"
PERIOD_END = "2026-05-15"

IB_CSV_PATH = "inbound.csv"
OB_CSV_PATH = "outbound.csv"
DEPT_CSV_PATH = "dept_calls.csv"
LOGO_PATH = "PromptPath_Logo.png"
OUTPUT_PATH = "PromptPath_Report.docx"

IB_ONLY = False
STORE_FILTER = None  # e.g. None | "All Star" | "All Star, Genesis Baton Rouge" | ["All Star", "Genesis Baton Rouge"]

# Column sets for validation (Streamlit / callers)
REQUIRED_INBOUND_COLUMNS = (
    "Dealerships",
    "Period",
    "Inbound Calls",
    "Connected",
    "Total Appts",
    "Hard Appt",
    "Delighted",
    "Disappointed",
)
REQUIRED_OUTBOUND_COLUMNS = (
    "Dealerships",
    "Period",
    "Outbound Dials",
    "Hard Appt",
    "Soft Appt",
)

# Outbound rows can label connect count as "Connected" or (common on newer exports) "Unique Connected".
OB_CONNECTED_COLUMN_CANDIDATES = (
    "Connected",
    "Unique Connected Outbound",
    "Unique Connected",
)

# Wide leaderboard "calls by department" export: inbound call counts per department (not Sales — SALES stays from IB CSV).
WIDE_DEPT_INBOUND_COLUMNS: Dict[str, str] = {
    "Service Inbound Calls": "service",
    "Parts Inbound Calls": "parts",
    "Finance Inbound Calls": "finance",
    "Other Inbound Calls": "other",
}

DEPT_LONG_CAT_MAP: Dict[str, str] = {
    "Sales Department": "sales",
    "Service Department": "service",
    "Parts Department": "parts",
    "Finance Department": "finance",
    "Other Department": "other",
}

# "Opportunities" rows: prefer unique-customer columns from leaderboard exports; else Connected.
IB_UNIQUE_OPP_COLUMNS = (
    "Unique Customer Inbound",
    "Unique Customers Inbound",
    "Unique Inbound Customers",
    "Unique Connected",
    "Unique Customer",
    "Unique Customers",
)
OB_UNIQUE_OPP_COLUMNS = (
    "Unique Customer Outbound",
    "Unique Customers Outbound",
    "Unique Outbound Customers",
    "Unique Connected",
    "Unique Customer",
    "Unique Customers",
)


def _first_matching_col(fieldnames: Optional[List[str]], candidates: tuple[str, ...]) -> Optional[str]:
    if not fieldnames:
        return None
    have = set(fieldnames)
    for c in candidates:
        if c in have:
            return c
    return None


def _resolve_ob_connected_column(fieldnames: Optional[List[str]]) -> Optional[str]:
    """Header for outbound connects (prefer legacy Connected when both exist)."""
    return _first_matching_col(fieldnames, OB_CONNECTED_COLUMN_CANDIDATES)


def _read_dept_csv(path: str) -> tuple[List[str], List[Dict[str, str]]]:
    with open(path, newline="", encoding="utf-8") as f:
        smp = f.read(1024)
        f.seek(0)
        dia = "excel-tab" if smp.count("\t") > smp.count(",") else "excel"
        r = csv.DictReader(f, dialect=dia)
        fn = list(r.fieldnames or [])
        return fn, list(r)


def _dept_csv_is_wide_format(fieldnames: List[str]) -> bool:
    fn = set(fieldnames)
    if "Dealerships" not in fn or "Period" not in fn:
        return False
    return all(col in fn for col in WIDE_DEPT_INBOUND_COLUMNS)


def _load_department_tables(path: str) -> tuple[Dict[str, Dict[str, int]], Dict[str, Dict[str, int]]]:
    """Return (current_period, previous_period) per-dealer dept call counts."""
    fn, rows = _read_dept_csv(path)
    dl_curr: Dict[str, Dict[str, int]] = {}
    dl_prev: Dict[str, Dict[str, int]] = {}

    if _dept_csv_is_wide_format(fn):
        for row in rows:
            dn = row.get("Dealerships", "").strip()
            if not dn or dn == "All Dealers":
                continue
            pe = row.get("Period", "").strip()
            if pe == "Current":
                target = dl_curr
            elif pe == "Previous":
                target = dl_prev
            else:
                continue
            if dn not in target:
                target[dn] = {}
            for col, key in WIDE_DEPT_INBOUND_COLUMNS.items():
                target[dn][key] = si(row.get(col, 0))
        return dl_curr, dl_prev

    has_period = "Period" in fn
    for row in rows:
        dn = row.get("dealer_name", row.get("Dealerships", "")).strip()
        cat = row.get("category", row.get("Category", "")).strip()
        calls = si(row.get("calls", row.get("Calls", 0)))
        key = DEPT_LONG_CAT_MAP.get(cat)
        if not dn or not key:
            continue
        if has_period:
            pe = row.get("Period", "").strip()
            if pe == "Current":
                target = dl_curr
            elif pe == "Previous":
                target = dl_prev
            else:
                continue
        else:
            target = dl_curr
        if dn not in target:
            target[dn] = {}
        target[dn][key] = calls
    return dl_curr, dl_prev


@dataclass
class ReportConfig:
    group_name: str
    period_start: str  # YYYY-MM-DD
    period_end: str
    ib_csv_path: str
    dept_csv_path: str
    logo_path: str
    output_path: str
    ob_csv_path: Optional[str] = None
    ib_only: bool = False
    # None = all stores. Otherwise OR-match: include store if name contains any substring.
    store_filter: Union[None, str, List[str]] = None
    # UI-selected line types for the listener sentence under Calls Processed by Department.
    listened_lines: List[str] = field(default_factory=list)
    # True for per-dealer DOCX from generate_dealer_reports (omit dept Group Total, etc.).
    single_store_report: bool = False


# Canonical labels for Streamlit multiselect and DOCX sentence (single source of truth).
LISTENED_LINE_TYPE_OPTIONS: tuple[str, ...] = (
    "Sales lines",
    "Operator lines",
    "Desk lines",
    "Service lines",
    "Parts lines",
)


def _lines_sentence(lines: Sequence[str]) -> str:
    known = set(LISTENED_LINE_TYPE_OPTIONS)
    sel = [x for x in lines if x in known]
    if not sel or set(sel) >= known:
        return "PromptPath currently listens to all your lines."
    labels = [x.lower() for x in sel]
    if len(labels) == 1:
        phrase = labels[0]
    else:
        phrase = ", ".join(labels[:-1]) + ", and " + labels[-1]
    return f"PromptPath currently listens to your {phrase}."


def normalize_store_filters(value: Union[None, str, Sequence[str]]) -> Optional[List[str]]:
    """Split optional filter input into non-empty substrings; None means no filter.

    Accepts a comma/semicolon/newline-separated string, a list/tuple of strings,
    or None. Matching is unchanged from v1: substring search, case-sensitive.
    """
    if value is None:
        return None
    if isinstance(value, str):
        parts = re.split(r"[,;\n]+", value)
        out = [p.strip() for p in parts if p.strip()]
        return out or None
    out = [str(x).strip() for x in value if str(x).strip()]
    return out or None


def store_matches_filters(
    store_name: str, filters: Optional[List[str]], *, exact: bool = False
) -> bool:
    if not filters:
        return True
    if exact:
        return store_name in filters
    return any(part in store_name for part in filters)


def _derive_report_strings(period_start: str, period_end: str) -> tuple[str, str]:
    s = datetime.strptime(period_start, "%Y-%m-%d")
    e = datetime.strptime(period_end, "%Y-%m-%d")
    report_period = (
        f"{s.strftime('%B %-d')} \u2013 {e.strftime('%-d, %Y')}"
        if s.month == e.month and s.year == e.year
        else f"{s.strftime('%B %-d')} \u2013 {e.strftime('%B %-d, %Y')}"
    )
    gen_date = e.strftime("%B %-d, %Y")
    return report_period, gen_date


def report_output_basename(dealer_or_group_name: str, period_start: str, period_end: str) -> str:
    """Filesystem-safe stem: {dealer}_{daterange}_promptpath_executive_report."""
    slug = re.sub(r"[^a-z0-9]+", "_", dealer_or_group_name.lower().strip()).strip("_") or "dealer"
    s = datetime.strptime(period_start, "%Y-%m-%d")
    e = datetime.strptime(period_end, "%Y-%m-%d")
    sm = s.strftime("%b").lower()
    em = e.strftime("%b").lower()
    if s.month == e.month and s.year == e.year:
        date_slug = f"{sm}{s.day}-{em}{e.day}_{s.year}"
    elif s.year == e.year:
        date_slug = f"{sm}{s.day}-{em}{e.day}_{s.year}"
    else:
        date_slug = f"{sm}{s.day}_{s.year}-{em}{e.day}_{e.year}"
    return f"{slug}_{date_slug}_promptpath_executive_report"


def _hex(c): return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

BRAND_NAVY = RGBColor(0x1D, 0x2D, 0x44)
BRAND_ORANGE = RGBColor(0xE0, 0x7B, 0x30)
LIGHT_BLUE = RGBColor(0xE8, 0xEE, 0xF7)
ORANGE_LIGHT = RGBColor(0xFD, 0xF0, 0xE4)
GREEN = RGBColor(0x1E, 0x7E, 0x4A)
RED = RGBColor(0xB9, 0x1C, 0x1C)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BORDER_GRAY = RGBColor(0xCC, 0xCC, 0xCC)
MID_GRAY = RGBColor(0x88, 0x88, 0x88)
IB_STRIPE = RGBColor(0xF0, 0xF4, 0xFA)
OB_STRIPE = RGBColor(0xFD, 0xF6, 0xEE)


def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for el in tcPr.findall(qn("w:shd")):
        tcPr.remove(el)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(color))
    tcPr.append(shd)


def set_cell_borders(cell, color=BORDER_GRAY, size=6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for el in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(el)
    tcB = OxmlElement("w:tcBorders")
    hx = _hex(color)
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), hx)
        tcB.append(b)
    tcPr.append(tcB)


def set_cell_margins(cell, top=60, bottom=60, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for el in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(el)
    tcM = OxmlElement("w:tcMar")
    for side, v in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(v))
        m.set(qn("w:type"), "dxa")
        tcM.append(m)
    tcPr.append(tcM)


def style_cell(cell, bg, bc=BORDER_GRAY, top=60, bottom=60, left=120, right=120):
    set_cell_bg(cell, bg)
    set_cell_borders(cell, bc)
    set_cell_margins(cell, top, bottom, left, right)


def cell_para(cell, text, bold=False, italic=False, size=9.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0] if cell.paragraphs and cell.paragraphs[0].text == "" else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Arial"
    return p


def add_run(para, text, bold=False, italic=False, size=10, color=GRAY):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Arial"
    return r


def set_col_widths(table, ws):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    for el in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(el)
    tblPr.append(lay)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), str(sum(ws)))
    tw.set(qn("w:type"), "dxa")
    for el in tblPr.findall(qn("w:tblW")):
        tblPr.remove(el)
    tblPr.append(tw)
    tg = tbl.find(qn("w:tblGrid"))
    if tg is None:
        tg = OxmlElement("w:tblGrid")
        tbl.insert(1, tg)
    else:
        for g in tg.findall(qn("w:gridCol")):
            tg.remove(g)
    for w in ws:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tg.append(gc)
    for row in table.rows:
        for ci, w in enumerate(ws):
            if ci < len(row.cells):
                tc = row.cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(w))
                tcW.set(qn("w:type"), "dxa")
                for el in tcPr.findall(qn("w:tcW")):
                    tcPr.remove(el)
                tcPr.append(tcW)


def _apply_cant_split(table) -> None:
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def section_heading(doc, text, sb=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(5)
    add_run(p, text, bold=True, size=13, color=BRAND_NAVY)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "10")
    bot.set(qn("w:space"), "3")
    bot.set(qn("w:color"), _hex(BRAND_ORANGE))
    pBdr.append(bot)
    pPr.append(pBdr)


def store_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, f"  {text}", bold=True, size=11, color=WHITE)
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:keepNext"))
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(BRAND_NAVY))
    pPr.append(shd)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "120")
    pPr.append(ind)


def rs(n):
    return {1: "1st", 2: "2nd", 3: "3rd"}.get(n, f"{n}th")


def rc(rank, total):
    return GREEN if rank == 1 else (RED if rank == total else GRAY)


def dstr(c, p):
    if p is None or p == 0:
        return "—", None
    chg = round((c - p) / p * 100)
    if chg == 0:
        return "—", None
    return f"{'▲' if chg > 0 else '▼'} {'+' if chg > 0 else ''}{chg}%", chg > 0


def dcol(s, hib=True):
    u = s.startswith("▲")
    d = s.startswith("▼")
    return (GREEN if u else RED if d else MID_GRAY) if hib else (RED if u else GREEN if d else MID_GRAY)


def _dept_count_cell(
    cell,
    bg,
    curr: int,
    prior: Optional[int],
    *,
    bold: bool = False,
    tc_=GRAY,
    align=WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    style_cell(cell, bg, BORDER_GRAY)
    disp = f"{curr:,}" if curr > 0 else "—"
    p = cell_para(cell, disp, bold=bold, size=9.5, color=tc_, align=align)
    chg, _ = dstr(curr, prior) if prior is not None else ("—", None)
    if chg != "—":
        add_run(p, f"\n{chg}", size=8, color=dcol(chg, hib=True))


def pct(n, d):
    return round(n / d * 100) if d else 0


def pct_appt_rate(total_appts: int, denom: int) -> float:
    """Appointment set rate: Total Appts ÷ denominator (unique customer opps), one decimal."""
    if not denom:
        return 0.0
    return round(100.0 * total_appts / denom, 1)


def fmt_appt_pct_label(v: float) -> str:
    """Display rate: integer percent when exact, else one decimal."""
    if abs(v - round(v)) < 1e-9:
        return f"{int(round(v))}%"
    return f"{v:.1f}%"


def pct_hard_of_hard_plus_soft(hard: int, soft: int) -> int:
    """Hard ÷ (Hard + Soft), whole percent — not ÷ Total Calls / Connected."""
    denom = hard + soft
    return round(100.0 * hard / denom) if denom else 0


def fmt_appts_hard_line(total_appts: int, hard_appts: int, soft_appts: int) -> str:
    hp = pct_hard_of_hard_plus_soft(hard_appts, soft_appts)
    return f"{total_appts} ({hp}% hard)"


def fs(c, calls):
    return f"{c} ({pct(c, calls)}%)"


def si(v):
    try:
        return int(float(str(v).strip()))
    except Exception:
        return 0


def _validate_csv_columns(fieldnames: Optional[List[str]], required: tuple, label: str) -> None:
    if not fieldnames:
        raise ValueError(f"{label}: file is empty or has no header row.")
    missing = [c for c in required if c not in fieldnames]
    if missing:
        raise ValueError(f"{label}: missing required column(s): {', '.join(missing)}.")


def validate_inbound_csv(path: str) -> None:
    with open(path, newline="", encoding="utf-8") as f:
        r = csv.DictReader(f)
        _validate_csv_columns(r.fieldnames, REQUIRED_INBOUND_COLUMNS, "Inbound CSV")


def validate_outbound_csv(path: str) -> None:
    with open(path, newline="", encoding="utf-8") as f:
        r = csv.DictReader(f)
        fn = r.fieldnames
    _validate_csv_columns(fn, REQUIRED_OUTBOUND_COLUMNS, "Outbound CSV")
    if _resolve_ob_connected_column(fn) is None:
        raise ValueError(
            "Outbound CSV: missing a Connected / connect-count column. "
            'Expected one of: "Connected", "Unique Connected Outbound", "Unique Connected".'
        )


def validate_dept_csv(path: str) -> None:
    fn, _ = _read_dept_csv(path)
    if not fn:
        raise ValueError("Department CSV: file is empty or has no header row.")
    if _dept_csv_is_wide_format(fn):
        missing = [c for c in WIDE_DEPT_INBOUND_COLUMNS if c not in fn]
        if missing:
            raise ValueError(
                "Department CSV (wide format): missing column(s): "
                + ", ".join(missing)
                + '. Expected Dealerships, Period, and the four "* Inbound Calls" department columns.'
            )
        return
    has_dealer = "dealer_name" in fn or "Dealerships" in fn
    has_cat = "category" in fn or "Category" in fn
    has_calls = "calls" in fn or "Calls" in fn
    if not (has_dealer and has_cat and has_calls):
        raise ValueError(
            "Department CSV: use either (1) wide leaderboard export with Dealerships, Period, and "
            "Service/Parts/Finance/Other Inbound Calls columns, or (2) long format with dealer, "
            "category (Sales/Service/Parts/Finance/Other Department), and calls columns."
        )


def generate_report(config: ReportConfig) -> str:
    """Build the DOCX at config.output_path and return that path."""
    report_period, gen_date = _derive_report_strings(config.period_start, config.period_end)
    group_name = config.group_name
    ib_only = config.ib_only

    with open(config.ib_csv_path, newline="", encoding="utf-8") as f:
        ib = list(csv.DictReader(f))

    inbound_soft_col = bool(ib and ib[0] is not None and "Soft Appt" in ib[0])
    _ib_fields = list(ib[0].keys()) if ib else []
    ib_opp_col = _first_matching_col(_ib_fields, IB_UNIQUE_OPP_COLUMNS) or "Connected"

    dd: Dict[str, Dict[str, Any]] = {}
    for row in ib:
        dn = row["Dealerships"].strip()
        pe = row["Period"].strip()
        if dn == "All Dealers":
            continue
        if dn not in dd:
            dd[dn] = {}
        d = dd[dn]
        px = "curr_" if pe == "Current" else "prev_"
        tot = si(row["Total Appts"])
        hard = si(row["Hard Appt"])
        if inbound_soft_col:
            raw_s = str(row.get("Soft Appt", "")).strip()
            soft = si(raw_s) if raw_s != "" else max(0, tot - hard)
        else:
            soft = max(0, tot - hard)
        d[px + "ib_calls"] = si(row["Inbound Calls"])
        d[px + "connected"] = si(row["Connected"])
        d[px + "ib_unique_opps"] = si(row[ib_opp_col])
        d[px + "total_appts"] = tot
        d[px + "hard_appts"] = hard
        d[px + "soft_appts"] = soft
        d[px + "delighted"] = si(row["Delighted"])
        d[px + "disappointed"] = si(row["Disappointed"])

    def pib(row):
        if not row:
            return {}
        tot = si(row["Total Appts"])
        hard = si(row["Hard Appt"])
        if inbound_soft_col:
            raw_s = str(row.get("Soft Appt", "")).strip()
            soft = si(raw_s) if raw_s != "" else max(0, tot - hard)
        else:
            soft = max(0, tot - hard)
        opp = si(row[ib_opp_col])
        return {
            k: si(row[v])
            for k, v in [
                ("ib_calls", "Inbound Calls"),
                ("connected", "Connected"),
                ("total_appts", "Total Appts"),
                ("hard_appts", "Hard Appt"),
                ("delighted", "Delighted"),
                ("disappointed", "Disappointed"),
            ]
        } | {"soft_appts": soft, "ib_unique_opps": opp}

    ac = pib(next((r for r in ib if r["Dealerships"] == "All Dealers" and r["Period"] == "Current"), None))
    ap = pib(next((r for r in ib if r["Dealerships"] == "All Dealers" and r["Period"] == "Previous"), None))
    for roll in (ac, ap):
        if roll:
            roll["curr_ib_unique_opps"] = roll.get("ib_unique_opps", 0)

    oc_d: Dict[str, Any] = {}
    op_d: Dict[str, Any] = {}
    ob: List[Dict[str, str]] = []
    ob_opp_col: Optional[str] = None
    ob_connected_col = "Connected"
    if not ib_only:
        if not config.ob_csv_path:
            raise ValueError("Outbound CSV path is required when ib_only is False.")
        with open(config.ob_csv_path, newline="", encoding="utf-8") as f:
            ob = list(csv.DictReader(f))
        _ob_fields = list(ob[0].keys()) if ob else []
        ob_connected_col = _resolve_ob_connected_column(_ob_fields) or "Connected"
        ob_opp_col = _first_matching_col(_ob_fields, OB_UNIQUE_OPP_COLUMNS) or ob_connected_col
        for row in ob:
            dn = row["Dealerships"].strip()
            pe = row["Period"].strip()
            if dn == "All Dealers":
                continue
            if dn not in dd:
                dd[dn] = {}
            d = dd[dn]
            px = "curr_" if pe == "Current" else "prev_"
            h_ob = si(row["Hard Appt"])
            s_ob = si(row["Soft Appt"])
            d[px + "ob_dials"] = si(row["Outbound Dials"])
            d[px + "ob_connected"] = si(row[ob_connected_col])
            d[px + "ob_unique_opps"] = si(row[ob_opp_col])
            d[px + "ob_total_appts"] = h_ob + s_ob
            d[px + "ob_hard_appts"] = h_ob
            d[px + "ob_soft_appts"] = s_ob

        def pob(row):
            if not row:
                return {}
            h_ob = si(row["Hard Appt"])
            s_ob = si(row["Soft Appt"])
            return {
                "ob_dials": si(row["Outbound Dials"]),
                "ob_connected": si(row[ob_connected_col]),
                "ob_unique_opps": si(row[ob_opp_col]),
                "ob_total_appts": h_ob + s_ob,
                "ob_hard_appts": h_ob,
                "ob_soft_appts": s_ob,
            }

        oc_d = pob(next((r for r in ob if r["Dealerships"] == "All Dealers" and r["Period"] == "Current"), None))
        op_d = pob(next((r for r in ob if r["Dealerships"] == "All Dealers" and r["Period"] == "Previous"), None))
        for roll in (oc_d, op_d):
            if roll:
                roll["curr_ob_unique_opps"] = roll.get("ob_unique_opps", 0)

    dl_curr, dl_prev = _load_department_tables(config.dept_csv_path)

    sf = normalize_store_filters(config.store_filter)
    exact_store = config.single_store_report
    sn = [s for s in dd if store_matches_filters(s, sf, exact=exact_store)]
    sn.sort(key=lambda s: dd[s].get("curr_ib_calls", 0), reverse=True)
    ts = len(sn)

    def build(mf, pf, df, hib=True):
        rows = []
        for name in sn:
            d = dd[name]
            cv = mf(d)
            pv = pf(d)
            chg, _ = dstr(cv, pv) if pv is not None else ("—", None)
            rows.append({"name": name, "mtd": df(d, "curr_"), "change": chg, "_sv": cv, "_hp": pv is not None, "hib": hib})
        avc = mf(ac) if ac else 0
        avp = pf(ap) if ap else None
        achg, _ = dstr(avc, avp) if avp else ("—", None)
        avg = {
            "name": "Group Avg",
            "mtd": df(ac, "curr_") if ac else "—",
            "change": achg,
            "_sv": avc,
            "rank": None,
            "hib": hib,
        }
        rows.sort(key=lambda x: x["_sv"], reverse=hib)
        for i, r in enumerate(rows):
            r["rank"] = i + 1
        res = []
        ins = False
        for r in rows:
            if not ins and (r["_sv"] <= avc if hib else r["_sv"] >= avc):
                res.append(avg)
                ins = True
            res.append(r)
        if not ins:
            res.append(avg)
        return res

    ibc = build(
        lambda d: d.get("curr_ib_unique_opps", 0),
        lambda d: d.get("prev_ib_unique_opps", None),
        lambda d, p: str(d.get(p + "ib_unique_opps", 0)),
    )
    iba = build(
        lambda d: d.get("curr_total_appts", 0),
        lambda d: d.get("prev_total_appts", None),
        lambda d, p: fmt_appts_hard_line(
            d.get(p + "total_appts", 0),
            d.get(p + "hard_appts", 0),
            d.get(p + "soft_appts", max(0, d.get(p + "total_appts", 0) - d.get(p + "hard_appts", 0))),
        ),
    )
    ibr = build(
        lambda d: pct_appt_rate(d.get("curr_total_appts", 0), d.get("curr_ib_unique_opps", 0)),
        lambda d: pct_appt_rate(d.get("prev_total_appts", 0), d.get("prev_ib_unique_opps", 0))
        if d.get("prev_ib_unique_opps")
        else None,
        lambda d, p: f"{pct_appt_rate(d.get(p + 'total_appts', 0), d.get(p + 'ib_unique_opps', 0))}%",
    )
    ibd = build(
        lambda d: pct(d.get("curr_delighted", 0), d.get("curr_ib_calls", 1)),
        lambda d: pct(d.get("prev_delighted", 0), d.get("prev_ib_calls", 1)) if d.get("prev_ib_calls") else None,
        lambda d, p: fs(d.get(p + "delighted", 0), d.get(p + "ib_calls", 1)),
    )
    ibx = build(
        lambda d: pct(d.get("curr_disappointed", 0), d.get("curr_ib_calls", 1)),
        lambda d: pct(d.get("prev_disappointed", 0), d.get("prev_ib_calls", 1)) if d.get("prev_ib_calls") else None,
        lambda d, p: fs(d.get(p + "disappointed", 0), d.get(p + "ib_calls", 1)),
        hib=False,
    )
    ibcr = build(
        lambda d: pct(d.get("curr_connected", 0), d.get("curr_ib_calls", 1)),
        lambda d: pct(d.get("prev_connected", 0), d.get("prev_ib_calls", 1)) if d.get("prev_ib_calls") else None,
        lambda d, p: f"{pct(d.get(p + 'connected', 0), d.get(p + 'ib_calls', 1))}%",
    )
    ib_b = [
        ("IB Connect Rate", ibcr, "inbound"),
        ("Opportunities", ibc, "inbound"),
        ("Appts Set (% hard)", iba, "inbound"),
        ("Appt Set Rate", ibr, "inbound"),
        ("Delighted Customers", ibd, "inbound"),
        ("Disappointed Customers", ibx, "inbound"),
    ]

    ob_b = []
    if not ib_only:
        ocr = build(
            lambda d: pct(d.get("curr_ob_connected", 0), d.get("curr_ob_dials", 1)),
            lambda d: pct(d.get("prev_ob_connected", 0), d.get("prev_ob_dials", 1)) if d.get("prev_ob_dials") else None,
            lambda d, p: f"{pct(d.get(p + 'ob_connected', 0), d.get(p + 'ob_dials', 1))}%",
        )
        oco = build(
            lambda d: d.get("curr_ob_unique_opps", 0),
            lambda d: d.get("prev_ob_unique_opps", None),
            lambda d, p: str(d.get(p + "ob_unique_opps", 0)),
        )
        oa = build(
            lambda d: d.get("curr_ob_total_appts", 0),
            lambda d: d.get("prev_ob_total_appts", None),
            lambda d, p: fmt_appts_hard_line(
                d.get(p + "ob_total_appts", 0),
                d.get(p + "ob_hard_appts", 0),
                d.get(p + "ob_soft_appts", max(0, d.get(p + "ob_total_appts", 0) - d.get(p + "ob_hard_appts", 0))),
            ),
        )
        or_ = build(
            lambda d: pct_appt_rate(d.get("curr_ob_total_appts", 0), d.get("curr_ob_connected", 0)),
            lambda d: pct_appt_rate(d.get("prev_ob_total_appts", 0), d.get("prev_ob_connected", 0))
            if d.get("prev_ob_connected")
            else None,
            lambda d, p: f"{pct_appt_rate(d.get(p + 'ob_total_appts', 0), d.get(p + 'ob_connected', 0))}%",
        )
        ob_b = [
            ("Connect Rate", ocr, "outbound"),
            ("Opportunities", oco, "outbound"),
            ("Appts Set (% hard)", oa, "outbound"),
            ("Appt Set Rate", or_, "outbound"),
        ]

    def footer(doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        tb = OxmlElement("w:top")
        tb.set(qn("w:val"), "single")
        tb.set(qn("w:sz"), "4")
        tb.set(qn("w:space"), "4")
        tb.set(qn("w:color"), "CCCCCC")
        pBdr.append(tb)
        pPr.append(pBdr)
        add_run(
            p,
            "Generated automatically by PromptPath. To adjust frequency or recipients, contact your Customer Success Manager.",
            size=8,
            color=MID_GRAY,
            italic=True,
        )
        fp = doc.sections[0].footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)
        fp.paragraph_format.space_after = Pt(4)

        def fld(para, code):
            r = para.add_run()
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "begin")
            r._r.append(fc)
            r2 = para.add_run()
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = code
            r2._r.append(it)
            r3 = para.add_run()
            fc2 = OxmlElement("w:fldChar")
            fc2.set(qn("w:fldCharType"), "end")
            r3._r.append(fc2)

        def fr(t):
            r = fp.add_run(t)
            r.font.size = Pt(8)
            r.font.color.rgb = MID_GRAY
            r.font.name = "Arial"
            r.font.italic = True

        fr("Page ")
        fld(fp, " PAGE ")
        fr(" of ")
        fld(fp, " NUMPAGES ")
        fr(f"  \u2022  {group_name}  \u2022  PromptPath  \u2022  \u00a9 2026 PromptPath. Proprietary and Confidential.")

    def make_header(doc):
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        for ci in range(2):
            c = t.rows[0].cells[ci]
            tc = c._tc
            tcPr = tc.get_or_add_tcPr()
            tb = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "none")
                tb.append(b)
            for el in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(el)
            tcPr.append(tb)
            set_cell_margins(c, top=0, bottom=0, left=0, right=0)
            set_cell_bg(c, WHITE)
        left = t.rows[0].cells[0]
        p1 = left.paragraphs[0]
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(3)
        add_run(p1, f"{group_name} \u2014 Dealer Performance Recap", bold=True, size=20, color=BRAND_NAVY)
        p2 = left.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        add_run(p2, report_period, bold=True, size=11, color=GRAY)
        p3 = left.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(0)
        add_run(p3, f"Generated {gen_date}  \u2022  MTD vs. Prior Month MTD", size=8.5, color=MID_GRAY, italic=True)
        pPr = p3._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "16")
        bot.set(qn("w:space"), "4")
        bot.set(qn("w:color"), _hex(BRAND_NAVY))
        pBdr.append(bot)
        pPr.append(pBdr)
        right = t.rows[0].cells[1]
        set_cell_margins(right, top=0, bottom=0, left=120, right=0)
        lp = right.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        lp.add_run().add_picture(config.logo_path, width=Inches(1.8))
        set_col_widths(t, [6400, 2960])

    def make_dept_table(doc):
        cw = ([2520] + [1368] * 5) if ib_only else ([2400] + [1080] * 5 + [1560])
        dc = (
            ["STORE", "SALES", "SERVICE", "PARTS", "FINANCE", "OTHER"]
            if ib_only
            else ["STORE", "SALES", "SERVICE", "PARTS", "FINANCE", "OTHER", "SALES DIALS"]
        )
        metric_keys = ["sales", "service", "parts", "finance", "other"] + ([] if ib_only else ["dials"])
        rd = []
        for name in sn:
            d = dd[name]
            dept = dl_curr.get(name, {})
            dept_p = dl_prev.get(name, {})
            r = {
                "name": name,
                "sales": d.get("curr_ib_calls", 0),
                "sales_prev": d.get("prev_ib_calls"),
                "service": dept.get("service", 0),
                "service_prev": dept_p.get("service"),
                "parts": dept.get("parts", 0),
                "parts_prev": dept_p.get("parts"),
                "finance": dept.get("finance", 0),
                "finance_prev": dept_p.get("finance"),
                "other": dept.get("other", 0),
                "other_prev": dept_p.get("other"),
            }
            if not ib_only:
                r["dials"] = d.get("curr_ob_dials", 0)
                r["dials_prev"] = d.get("prev_ob_dials")
            rd.append(r)
        gt = {
            "name": "Group Total",
            "is_total": True,
            "sales": sum(r["sales"] for r in rd),
            "sales_prev": sum(dd[n].get("prev_ib_calls", 0) for n in sn),
            "service": sum(r["service"] for r in rd),
            "service_prev": sum(dl_prev.get(n, {}).get("service", 0) for n in sn),
            "parts": sum(r["parts"] for r in rd),
            "parts_prev": sum(dl_prev.get(n, {}).get("parts", 0) for n in sn),
            "finance": sum(r["finance"] for r in rd),
            "finance_prev": sum(dl_prev.get(n, {}).get("finance", 0) for n in sn),
            "other": sum(r["other"] for r in rd),
            "other_prev": sum(dl_prev.get(n, {}).get("other", 0) for n in sn),
        }
        if not ib_only:
            gt["dials"] = oc_d.get("ob_dials", 0)
            gt["dials_prev"] = sum(dd[n].get("prev_ob_dials", 0) for n in sn)
        ar = rd + ([gt] if len(rd) > 1 and not config.single_store_report else [])
        nc = len(dc)
        t = doc.add_table(rows=2 + len(ar), cols=nc)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        sup = t.rows[0]
        style_cell(sup.cells[0], WHITE, BORDER_GRAY, top=50, bottom=50)
        cell_para(sup.cells[0], "", size=8)
        m = sup.cells[1]
        for _ in range(2, 6):
            m = m.merge(sup.cells[_])
        style_cell(m, LIGHT_BLUE, BORDER_GRAY, top=55, bottom=55)
        cell_para(m, "INBOUND", bold=True, size=9, color=BRAND_NAVY)
        if not ib_only:
            style_cell(sup.cells[6], ORANGE_LIGHT, BORDER_GRAY, top=55, bottom=55)
            cell_para(sup.cells[6], "OUTBOUND", bold=True, size=9, color=BRAND_ORANGE)
        hdr = t.rows[1]
        bgs = [BRAND_NAVY] + [LIGHT_BLUE] * 5 + ([] if ib_only else [ORANGE_LIGHT])
        clrs = [WHITE] + [BRAND_NAVY] * 5 + ([] if ib_only else [BRAND_ORANGE])
        for ci, (col, bg, clr) in enumerate(zip(dc, bgs, clrs)):
            c = hdr.cells[ci]
            style_cell(c, bg, BORDER_GRAY, top=65, bottom=65)
            cell_para(
                c,
                col,
                bold=True,
                size=8.5,
                color=clr,
                align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER,
            )
        keys = ["name"] + metric_keys
        for ri, row_d in enumerate(ar):
            it = row_d.get("is_total", False)
            bg = BRAND_NAVY if it else (WHITE if ri % 2 == 0 else LIGHT_GRAY)
            tc_ = WHITE if it else GRAY
            row = t.rows[ri + 2]
            for ci, key in enumerate(keys):
                c = row.cells[ci]
                if ci == 0:
                    style_cell(c, bg, BORDER_GRAY)
                    cell_para(
                        c,
                        row_d["name"],
                        bold=it,
                        size=9.5,
                        color=tc_,
                        align=WD_ALIGN_PARAGRAPH.LEFT,
                    )
                else:
                    prior_val = row_d.get(f"{key}_prev")
                    _dept_count_cell(
                        c,
                        bg,
                        int(row_d[key]),
                        int(prior_val) if prior_val is not None else None,
                        bold=it,
                        tc_=tc_,
                    )
        set_col_widths(t, cw)
        _apply_cant_split(t)

    def kgp(doc):
        section_heading(doc, "Key Group Performance", sb=14)
        iu = ac.get("ib_unique_opps", 0)
        ia = ac.get("total_appts", 0)
        ir = pct_appt_rate(ia, iu)
        pr = pct_appt_rate(ap.get("total_appts", 0), ap.get("ib_unique_opps", 0))
        ic_ = round(ir - pr, 1)
        ia_ = f"{'▲' if ic_ > 0 else '▼' if ic_ < 0 else '—'} {'+' if ic_ > 0 else ''}{ic_}pp" if ic_ != 0 else "—"
        nr = 2 if ib_only else 3
        t = doc.add_table(rows=nr, cols=3)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for ci, h in enumerate(["METRIC", "MTD RATE", "CHANGE VS. PRIOR MTD"]):
            c = t.rows[0].cells[ci]
            style_cell(c, BRAND_NAVY, BORDER_GRAY, top=65, bottom=65)
            cell_para(
                c,
                h,
                bold=True,
                size=8.5,
                color=WHITE,
                align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER,
            )
        row = t.rows[1]
        lbl = f"IB Appt Set Rate  ({ia:,} appts / {iu:,} unique)"
        style_cell(row.cells[0], LIGHT_BLUE, BORDER_GRAY)
        cell_para(row.cells[0], lbl, bold=True, size=9.5, color=BRAND_NAVY, align=WD_ALIGN_PARAGRAPH.LEFT)
        style_cell(row.cells[1], LIGHT_BLUE, BORDER_GRAY)
        cell_para(row.cells[1], fmt_appt_pct_label(ir), bold=True, size=11, color=BRAND_NAVY)
        style_cell(row.cells[2], LIGHT_BLUE, BORDER_GRAY)
        cell_para(row.cells[2], ia_, bold=True, size=10, color=GREEN if ic_ > 0 else (RED if ic_ < 0 else MID_GRAY))
        if not ib_only:
            ou = oc_d.get("ob_connected", 0)
            ota = oc_d.get("ob_total_appts", 0)
            otr = pct_appt_rate(ota, ou)
            opr = pct_appt_rate(op_d.get("ob_total_appts", 0), op_d.get("ob_connected", 0))
            oc_ = round(otr - opr, 1)
            oa_ = f"{'▲' if oc_ > 0 else '▼' if oc_ < 0 else '—'} {'+' if oc_ > 0 else ''}{oc_}pp" if oc_ != 0 else "—"
            row = t.rows[2]
            olbl = f"OB Appt Set Rate  ({ota:,} appts / {ou:,} connected)"
            style_cell(row.cells[0], ORANGE_LIGHT, BORDER_GRAY)
            cell_para(row.cells[0], olbl, bold=True, size=9.5, color=BRAND_ORANGE, align=WD_ALIGN_PARAGRAPH.LEFT)
            style_cell(row.cells[1], ORANGE_LIGHT, BORDER_GRAY)
            cell_para(row.cells[1], fmt_appt_pct_label(otr), bold=True, size=11, color=BRAND_ORANGE)
            style_cell(row.cells[2], ORANGE_LIGHT, BORDER_GRAY)
            cell_para(row.cells[2], oa_, bold=True, size=10, color=GREEN if oc_ > 0 else (RED if oc_ < 0 else MID_GRAY))
        _apply_cant_split(t)
        set_col_widths(t, [4320, 2520, 2520])

    def store_bundle(doc, name):
        store_heading(doc, name.upper())
        mr = [(l, "inbound", next((r for r in b if r["name"] == name), None)) for l, b, _ in ib_b]
        mr += [(l, "outbound", next((r for r in b if r["name"] == name), None)) for l, b, _ in ob_b]
        mr = [(l, c, r) for l, c, r in mr if r]
        show_rank = ts > 1
        headers = ["", "METRIC", "MTD", "CHANGE"] + (["RANK"] if show_rank else [])
        alignments = (
            [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
            + [WD_ALIGN_PARAGRAPH.CENTER] * (2 + (1 if show_rank else 0))
        )
        t = doc.add_table(rows=1 + len(mr), cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for ci, (h, al) in enumerate(zip(headers, alignments)):
            c = t.rows[0].cells[ci]
            style_cell(c, BRAND_NAVY, BRAND_NAVY, top=65, bottom=65)
            cell_para(c, h, bold=True, size=8.5, color=WHITE, align=al)
        for ri, (lbl, ch, sr) in enumerate(mr):
            ib_row = ch == "inbound"
            mb = IB_STRIPE if ib_row else OB_STRIPE
            hb = LIGHT_BLUE if ib_row else ORANGE_LIGHT
            cc = BRAND_NAVY if ib_row else BRAND_ORANGE
            row = t.rows[1 + ri]
            style_cell(row.cells[0], hb, BORDER_GRAY, top=55, bottom=55)
            cell_para(row.cells[0], "I" if ib_row else "O", bold=True, size=8, color=cc)
            style_cell(row.cells[1], mb, BORDER_GRAY, top=55, bottom=55)
            cell_para(row.cells[1], lbl, bold=True, size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.LEFT)
            style_cell(row.cells[2], mb, BORDER_GRAY, top=55, bottom=55)
            cell_para(row.cells[2], sr["mtd"], bold=True, size=9.5, color=GRAY)
            style_cell(row.cells[3], mb, BORDER_GRAY, top=55, bottom=55)
            chg = sr["change"]
            hib = sr.get("hib", True)
            cell_para(row.cells[3], chg, size=9.5, color=dcol(chg, hib) if chg != "—" else MID_GRAY)
            if show_rank:
                style_cell(row.cells[4], mb, BORDER_GRAY, top=55, bottom=55)
                rank = sr.get("rank")
                cell_para(
                    row.cells[4],
                    f"{rs(rank)} of {ts}" if rank else "—",
                    bold=bool(rank),
                    size=9,
                    color=rc(rank, ts) if rank else MID_GRAY,
                )
        _apply_cant_split(t)
        set_col_widths(
            t,
            [400, 5040, 1800, 1440, 1160] if show_rank else [400, 6200, 1800, 1440],
        )

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    make_header(doc)
    section_heading(doc, "Calls Processed by Department", sb=12)
    p_lines = doc.add_paragraph()
    p_lines.paragraph_format.space_before = Pt(0)
    p_lines.paragraph_format.space_after = Pt(6)
    add_run(p_lines, _lines_sentence(config.listened_lines), italic=True, size=9, color=MID_GRAY)
    make_dept_table(doc)
    kgp(doc)
    section_heading(doc, "Store Sales Performance Summaries", sb=14)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    hn = any(not r.get("_hp", True) for _, b, _ in ib_b for r in b if r["name"] != "Group Avg")
    note = (
        "Inbound Opportunities use unique customer counts from the CSV when available (otherwise Connected); "
        "IB appointment set rate uses that inbound unique count as its denominator."
    )
    if not ib_only:
        note += (
            "  Outbound Opportunities still use unique counts when available (otherwise Connected); "
            "OB appointment set rate divides outbound appointments by outbound Connected."
        )
    if hn:
        note += "  Stores showing \u2014 in Change have no prior month data."
    add_run(p, note, italic=True, size=9, color=MID_GRAY)
    for name in sn:
        store_bundle(doc, name)
    footer(doc)
    doc.save(config.output_path)
    audit_path = str(Path(config.output_path).with_name(Path(config.output_path).stem + "_audit.xlsx"))
    from report_audit import write_audit_workbook

    write_audit_workbook(
        output_xlsx_path=audit_path,
        group_name=group_name,
        period_label=f"{report_period} | Generated {gen_date}",
        ib_rows=ib,
        ob_rows=ob if ob else None,
        ib_only=ib_only,
        sn_ordered=list(sn),
        dd=dd,
        inbound_has_soft_appt_column=inbound_soft_col,
        ib_opportunities_column=ib_opp_col,
        ob_opportunities_column=ob_opp_col,
        ob_connected_column=ob_connected_col,
    )
    return config.output_path


def _filtered_sorted_store_names(config: ReportConfig) -> List[str]:
    """Same dealer ordering as generate_report (filters + sort by current inbound calls)."""
    ib_only = config.ib_only
    with open(config.ib_csv_path, newline="", encoding="utf-8") as f:
        ib = list(csv.DictReader(f))

    inbound_soft_col = bool(ib and ib[0] is not None and "Soft Appt" in ib[0])
    _ib_fields = list(ib[0].keys()) if ib else []
    ib_opp_col = _first_matching_col(_ib_fields, IB_UNIQUE_OPP_COLUMNS) or "Connected"

    dd: Dict[str, Dict[str, Any]] = {}
    for row in ib:
        dn = row["Dealerships"].strip()
        pe = row["Period"].strip()
        if dn == "All Dealers":
            continue
        if dn not in dd:
            dd[dn] = {}
        d = dd[dn]
        px = "curr_" if pe == "Current" else "prev_"
        tot = si(row["Total Appts"])
        hard = si(row["Hard Appt"])
        if inbound_soft_col:
            raw_s = str(row.get("Soft Appt", "")).strip()
            soft = si(raw_s) if raw_s != "" else max(0, tot - hard)
        else:
            soft = max(0, tot - hard)
        d[px + "ib_calls"] = si(row["Inbound Calls"])
        d[px + "connected"] = si(row["Connected"])
        d[px + "ib_unique_opps"] = si(row[ib_opp_col])
        d[px + "total_appts"] = tot
        d[px + "hard_appts"] = hard
        d[px + "soft_appts"] = soft
        d[px + "delighted"] = si(row["Delighted"])
        d[px + "disappointed"] = si(row["Disappointed"])

    if not ib_only:
        if not config.ob_csv_path:
            raise ValueError("Outbound CSV path is required when ib_only is False.")
        with open(config.ob_csv_path, newline="", encoding="utf-8") as f:
            ob = list(csv.DictReader(f))
        _ob_fields = list(ob[0].keys()) if ob else []
        ob_connected_col = _resolve_ob_connected_column(_ob_fields) or "Connected"
        ob_opp_col = _first_matching_col(_ob_fields, OB_UNIQUE_OPP_COLUMNS) or ob_connected_col
        for row in ob:
            dn = row["Dealerships"].strip()
            pe = row["Period"].strip()
            if dn == "All Dealers":
                continue
            if dn not in dd:
                dd[dn] = {}
            d = dd[dn]
            px = "curr_" if pe == "Current" else "prev_"
            h_ob = si(row["Hard Appt"])
            s_ob = si(row["Soft Appt"])
            d[px + "ob_dials"] = si(row["Outbound Dials"])
            d[px + "ob_connected"] = si(row[ob_connected_col])
            d[px + "ob_unique_opps"] = si(row[ob_opp_col])
            d[px + "ob_total_appts"] = h_ob + s_ob
            d[px + "ob_hard_appts"] = h_ob
            d[px + "ob_soft_appts"] = s_ob

    sf = normalize_store_filters(config.store_filter)
    sn = [s for s in dd if store_matches_filters(s, sf)]
    sn.sort(key=lambda s: dd[s].get("curr_ib_calls", 0), reverse=True)
    return sn


def generate_dealer_reports(config: ReportConfig, output_dir: str) -> List[str]:
    """Generate one full exec-style DOCX per dealer in config's filtered store list.

    Writes companion *_audit.xlsx next to each DOCX via generate_report.
    Returns paths in the same order as the executive report store list.
    """
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    dealers = _filtered_sorted_store_names(config)
    paths: List[str] = []
    for dealer in dealers:
        stem = report_output_basename(dealer, config.period_start, config.period_end)
        out_path = str(Path(output_dir) / f"{stem}.docx")
        dealer_cfg = replace(
            config,
            store_filter=dealer,
            output_path=out_path,
            single_store_report=True,
        )
        generate_report(dealer_cfg)
        paths.append(out_path)
    return paths


def _main_cli():
    cfg = ReportConfig(
        group_name=GROUP_NAME,
        period_start=PERIOD_START,
        period_end=PERIOD_END,
        ib_csv_path=IB_CSV_PATH,
        ob_csv_path=None if IB_ONLY else OB_CSV_PATH,
        dept_csv_path=DEPT_CSV_PATH,
        logo_path=LOGO_PATH,
        output_path=OUTPUT_PATH,
        ib_only=IB_ONLY,
        store_filter=STORE_FILTER,
    )
    path = generate_report(cfg)
    audit_xlsx = str(Path(path).with_name(Path(path).stem + "_audit.xlsx"))
    report_period, _ = _derive_report_strings(cfg.period_start, cfg.period_end)
    with open(cfg.ib_csv_path, newline="", encoding="utf-8") as f:
        ib = list(csv.DictReader(f))
    dd = {}
    for row in ib:
        dn = row["Dealerships"].strip()
        pe = row["Period"].strip()
        if dn == "All Dealers":
            continue
        if dn not in dd:
            dd[dn] = {}
        d = dd[dn]
        px = "curr_" if pe == "Current" else "prev_"
        d[px + "ib_calls"] = si(row["Inbound Calls"])
        d[px + "connected"] = si(row["Connected"])
        d[px + "total_appts"] = si(row["Total Appts"])
        d[px + "hard_appts"] = si(row["Hard Appt"])
        d[px + "delighted"] = si(row["Delighted"])
        d[px + "disappointed"] = si(row["Disappointed"])
    sf_cli = normalize_store_filters(cfg.store_filter)
    sn = [s for s in dd if store_matches_filters(s, sf_cli)]
    sn.sort(key=lambda s: dd[s].get("curr_ib_calls", 0), reverse=True)
    ts = len(sn)
    filter_disp = ", ".join(sf_cli) if sf_cli else "none"
    print(f"\n\u2713  Saved: {path}")
    if Path(audit_xlsx).is_file():
        print(f"   Audit workbook: {audit_xlsx}")
    print(f"   Group: {cfg.group_name}  |  Period: {report_period}  |  Stores: {ts}")
    print(f"   IB only: {cfg.ib_only}  |  Filter: {filter_disp}")
    for s in sn:
        d = dd[s]
        hp = "prev_connected" in d
        print(f"   {'✓' if hp else '—'}  {s}  ({d.get('curr_ib_calls', 0)} inbound calls)")


if __name__ == "__main__":
    _main_cli()
