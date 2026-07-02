"""
PromptCast script generation for PromptPath Executive Reports.

Builds a data brief from already-computed ReportData, extracts the text of the
generated executive report DOCX, then calls OpenAI to produce a ~300-word spoken-
prose audio script following the Marcus A. voice prompt.

The script is written to a branded .docx ready for download and pasting into Hume.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from promptpath_exec_report_v1 import (
    ReportConfig,
    ReportData,
    _derive_report_strings,
    derive_comparison_period_label,
    pct_appt_rate,
    pct,
)

# ── constants ────────────────────────────────────────────────────────────────

PLATFORM_BENCHMARK_LABEL = "40 to 45 percent inbound appointment set rate"

_PROMPT_FILE = Path(__file__).resolve().parent / "promptcast-prompt-marcus.txt"

_BRAND_NAVY = RGBColor(0x1D, 0x2D, 0x44)
_BRAND_ORANGE = RGBColor(0xE0, 0x7B, 0x30)
_GRAY = RGBColor(0x55, 0x55, 0x55)
_MID_GRAY = RGBColor(0x88, 0x88, 0x88)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Abbreviations that TTS commonly mispronounces; add more as needed.
_TTS_EXPANSIONS: Dict[str, str] = {
    "BMW": "B M W",
    "CDJR": "C D J R",
    "CDJ": "C D J",
    "GMC": "G M C",
    "INFINITI": "Infiniti",
    "KIA": "K I A",
    "VW": "V W",
}


# ── DOCX text extraction ─────────────────────────────────────────────────────

def extract_docx_text(docx_path: str) -> str:
    """Return a plain-text transcript of the executive report DOCX.

    Paragraphs are rendered as lines; tables are rendered as pipe-separated rows.
    Header/footer runs are intentionally omitted.
    """
    doc = Document(docx_path)
    lines: List[str] = []

    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            from docx.text.paragraph import Paragraph  # type: ignore
            para = Paragraph(child, doc)
            text = para.text.strip()
            if text:
                lines.append(text)
        elif tag == "tbl":
            from docx.table import Table  # type: ignore
            tbl = Table(child, doc)
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                # Deduplicate merged-cell repetitions
                deduped: List[str] = []
                for c in cells:
                    if not deduped or c != deduped[-1]:
                        deduped.append(c)
                row_text = " | ".join(deduped)
                if row_text.strip(" |"):
                    lines.append(row_text)

    return "\n".join(lines)


# ── data brief builder ───────────────────────────────────────────────────────

def build_data_brief(
    config: ReportConfig,
    data: ReportData,
    ramp_status: str,
    outbound_launch_date: str,
    report_period: str,
    comparison_period: str,
) -> str:
    """Build the structured data brief sent to OpenAI as the user message (part 1)."""
    dd = data.dd
    ac = data.ac
    ap = data.ap
    oc_d = data.oc_d
    op_d = data.op_d
    dl_curr = data.dl_curr
    sn = data.sn
    ib_only = config.ib_only

    lines: List[str] = []
    a = lines.append

    # ── header fields ────────────────────────────────────────────────────────
    a("=== GROUP / STORE INFORMATION ===")
    a(f"Group / Store Name: {config.group_name}")
    a("Report Type: Group Level")
    a(f"Report Period: {report_period}")
    a(f"Comparison Period: {comparison_period}")
    a(f"Number of Stores: {len(sn)}")
    a(f"Ramp Status: {ramp_status}")
    if not ib_only and outbound_launch_date.strip():
        a(f"Outbound Launch Date: {outbound_launch_date.strip()}")
    a(f"Platform Benchmark: {PLATFORM_BENCHMARK_LABEL}")
    a("")

    # ── department call volumes (Call Insights section) ──────────────────────
    a("=== CALL VOLUME BY DEPARTMENT (GROUP TOTALS) ===")
    sales_curr = sum(dd[n].get("curr_ib_calls", 0) for n in sn)
    service_curr = sum(dl_curr.get(n, {}).get("service", 0) for n in sn)
    parts_curr = sum(dl_curr.get(n, {}).get("parts", 0) for n in sn)
    finance_curr = sum(dl_curr.get(n, {}).get("finance", 0) for n in sn)
    other_curr = sum(dl_curr.get(n, {}).get("other", 0) for n in sn)
    total_ib_curr = sales_curr + service_curr + parts_curr + finance_curr + other_curr
    a(f"Sales Inbound Calls (current): {sales_curr:,}")
    a(f"Service Inbound Calls (current): {service_curr:,}")
    a(f"Parts Inbound Calls (current): {parts_curr:,}")
    a(f"Finance Inbound Calls (current): {finance_curr:,}")
    a(f"Other Inbound Calls (current): {other_curr:,}")
    a(f"Total Inbound Calls (current): {total_ib_curr:,}")
    if not ib_only:
        total_dials_curr = sum(dd[n].get("curr_ob_dials", 0) for n in sn)
        a(f"Total Outbound Dials (current): {total_dials_curr:,}")
    a("")

    # ── key group performance (IB appt set rate) ─────────────────────────────
    a("=== KEY GROUP PERFORMANCE ===")
    iu_c = ac.get("ib_unique_opps", 0)
    ia_c = ac.get("total_appts", 0)
    ib_asr_curr = pct_appt_rate(ia_c, iu_c)
    a(f"Group IB Appt Set Rate (current): {ib_asr_curr}%  ({ia_c:,} appts / {iu_c:,} unique opps)")
    iu_p = ap.get("ib_unique_opps", 0) if ap else 0
    ia_p = ap.get("total_appts", 0) if ap else 0
    ib_asr_prev = pct_appt_rate(ia_p, iu_p)
    a(f"Group IB Appt Set Rate (comparison): {ib_asr_prev}%  ({ia_p:,} appts / {iu_p:,} unique opps)")
    if not ib_only and oc_d:
        ob_conn_c = oc_d.get("ob_connected", 0)
        ob_apts_c = oc_d.get("ob_total_appts", 0)
        ob_asr_curr = pct_appt_rate(ob_apts_c, ob_conn_c)
        a(f"Group OB Appt Set Rate (current): {ob_asr_curr}%  ({ob_apts_c:,} appts / {ob_conn_c:,} connected)")
        if op_d:
            ob_conn_p = op_d.get("ob_connected", 0)
            ob_apts_p = op_d.get("ob_total_appts", 0)
            ob_asr_prev = pct_appt_rate(ob_apts_p, ob_conn_p)
            a(f"Group OB Appt Set Rate (comparison): {ob_asr_prev}%")
    a("")

    # ── per-store inbound table (sorted by current appt set rate desc) ───────
    a("=== PER-STORE INBOUND PERFORMANCE ===")
    a("Store | IB Appt Set Rate (curr) | IB Appt Set Rate (prev) | "
      "Connected Calls (curr) | Delighted % (curr) | Disappointed % (curr)")
    store_rows = []
    for name in sn:
        d = dd[name]
        iu = d.get("curr_ib_unique_opps", 0)
        ia = d.get("curr_total_appts", 0)
        asr_c = pct_appt_rate(ia, iu)
        iu_p2 = d.get("prev_ib_unique_opps", 0)
        ia_p2 = d.get("prev_total_appts", 0)
        asr_p2 = pct_appt_rate(ia_p2, iu_p2) if iu_p2 else None
        conn = d.get("curr_connected", 0)
        ib_calls = d.get("curr_ib_calls", 1)
        del_pct = pct(d.get("curr_delighted", 0), ib_calls)
        dis_pct = pct(d.get("curr_disappointed", 0), ib_calls)
        store_rows.append((asr_c, name, asr_p2, conn, del_pct, dis_pct))
    store_rows.sort(key=lambda x: x[0], reverse=True)
    for asr_c, name, asr_p2, conn, del_pct, dis_pct in store_rows:
        prev_str = f"{asr_p2}%" if asr_p2 is not None else "n/a"
        a(f"{name} | {asr_c}% | {prev_str} | {conn:,} connected | {del_pct}% delighted | {dis_pct}% disappointed")
    a("")

    # ── per-store outbound table ──────────────────────────────────────────────
    if not ib_only:
        a("=== PER-STORE OUTBOUND PERFORMANCE ===")
        a("Store | Dials (curr) | Dials (prev) | OB Appt Set Rate (curr) | OB Appt Set Rate (prev)")
        ob_rows = []
        for name in sn:
            d = dd[name]
            dials_c = d.get("curr_ob_dials", 0)
            dials_p = d.get("prev_ob_dials", None)
            ob_conn_c = d.get("curr_ob_connected", 0)
            ob_apts_c = d.get("curr_ob_total_appts", 0)
            ob_asr_c = pct_appt_rate(ob_apts_c, ob_conn_c)
            ob_conn_p = d.get("prev_ob_connected", 0)
            ob_apts_p = d.get("prev_ob_total_appts", 0)
            ob_asr_p = pct_appt_rate(ob_apts_p, ob_conn_p) if ob_conn_p else None
            ob_rows.append((dials_c, name, dials_p, ob_asr_c, ob_asr_p))
        ob_rows.sort(key=lambda x: x[0], reverse=True)
        for dials_c, name, dials_p, ob_asr_c, ob_asr_p in ob_rows:
            prev_dials = f"{dials_p:,}" if dials_p is not None else "n/a"
            prev_asr = f"{ob_asr_p}%" if ob_asr_p is not None else "n/a"
            a(f"{name} | {dials_c:,} | {prev_dials} | {ob_asr_c}% | {prev_asr}")
        a("")

    return "\n".join(lines)


# ── pronunciation helper ──────────────────────────────────────────────────────

def _expand_tts_abbreviations(text: str) -> str:
    """Replace known TTS-unfriendly abbreviations with their spoken expansions."""
    for abbr, expansion in _TTS_EXPANSIONS.items():
        text = re.sub(r"\b" + re.escape(abbr) + r"\b", expansion, text)
    return text


# ── user message builder ─────────────────────────────────────────────────────

def build_user_message(data_brief: str, report_text: str) -> str:
    """Combine the data brief, executive report text, and guardrail rules."""
    guardrails = (
        "=== GUARDRAIL RULES ===\n"
        "Only use figures that appear in the DATA BRIEF above or the EXECUTIVE REPORT below. "
        "Do not calculate, estimate, or invent any number that appears in neither source. "
        "If a metric is not present in either source, omit that section silently.\n"
        "For TTS pronunciation: spell out abbreviations that text-to-speech engines "
        "commonly mispronounce. For example: 'BMW' → 'B M W', 'CDJR' → 'C D J R'. "
        "Apply this rule to all dealership brand abbreviations in the script output."
    )
    return (
        f"=== DATA BRIEF ===\n{data_brief}\n\n"
        f"=== EXECUTIVE REPORT (verbatim) ===\n{report_text}\n\n"
        f"{guardrails}"
    )


# ── OpenAI call ───────────────────────────────────────────────────────────────

def generate_script(api_key: str, user_message: str, model: str = "gpt-4o") -> str:
    """Call OpenAI and return the plain spoken-prose PromptCast script."""
    from openai import OpenAI  # imported here so the module loads without openai installed

    system_prompt = _PROMPT_FILE.read_text(encoding="utf-8")
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
        temperature=0.4,
    )
    script = response.choices[0].message.content or ""
    return _expand_tts_abbreviations(script.strip())


# ── DOCX writer ───────────────────────────────────────────────────────────────

def write_promptcast_docx(
    script_text: str,
    output_path: str,
    group_name: str,
    period_label: str,
    logo_path: str,
) -> str:
    """Write a branded .docx containing the spoken-prose PromptCast script."""
    doc = Document()

    # Remove default section margins to match executive report style
    section = doc.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    # Header table: title left, logo right
    ht = doc.add_table(rows=1, cols=2)
    ht.style = "Table Grid"
    ht.autofit = False
    left = ht.rows[0].cells[0]
    right = ht.rows[0].cells[1]

    # Remove table borders
    for cell in (left, right):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    p_title = left.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_title.add_run("PromptCast Script")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = _BRAND_NAVY

    p_sub = left.add_paragraph()
    run2 = p_sub.add_run(f"{group_name}  —  {period_label}")
    run2.font.size = Pt(11)
    run2.font.color.rgb = _GRAY

    p_note = left.add_paragraph()
    run3 = p_note.add_run(
        "PromptCast audio script — generated by PromptPath. "
        "Paste into Hume for text-to-speech conversion."
    )
    run3.italic = True
    run3.font.size = Pt(8.5)
    run3.font.color.rgb = _MID_GRAY

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        rp.add_run().add_picture(logo_path, width=Inches(1.8))
    except Exception:
        pass

    # Set column widths
    from docx.oxml import OxmlElement as OE
    tbl_el = ht._tbl
    tblGrid = OE("w:tblGrid")
    for w in [6400, 2960]:
        gridCol = OE("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        tblGrid.append(gridCol)
    tbl_el.insert(0, tblGrid)

    # Separator line
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(12)
    sep.paragraph_format.space_after = Pt(12)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "16")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "1D2D44")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Script body — one paragraph of flowing prose
    p_script = doc.add_paragraph()
    p_script.paragraph_format.space_before = Pt(12)
    p_script.paragraph_format.space_after = Pt(0)
    run_script = p_script.add_run(script_text)
    run_script.font.size = Pt(11)
    run_script.font.color.rgb = _GRAY

    doc.save(output_path)
    return output_path
